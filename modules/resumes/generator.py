'''
Author:     Suraj Panwar
LinkedIn:   https://www.linkedin.com/in/surajpanwar/

Copyright (C) 2024-2026 Suraj Panwar

License:    GNU Affero General Public License
            https://www.gnu.org/licenses/agpl-3.0.en.html
            
GitHub:     https://github.com/surajpanwar/Auto_job_applier_linkedIn

Modified by: Suraj Panwar

'''

import os
from datetime import datetime
import docx
from fpdf import FPDF


def tailor_resume_summary(original_summary: str, job_description: str, 
                          required_skills: list = None) -> str:
    """
    Tailors a resume summary based on job description.
    
    Args:
        original_summary: The original professional summary
        job_description: The job description to tailor for
        required_skills: List of required skills from the job
        
    Returns:
        Tailored summary string
    """
    if required_skills and len(required_skills) > 0:
        skill_mention = f" Experienced in {', '.join(required_skills[:3])}."
        return original_summary + skill_mention
    return original_summary


def highlight_matching_skills(user_skills: list, job_skills: list) -> list:
    """
    Returns skills that match between user's skills and job requirements.
    
    Args:
        user_skills: List of user's skills
        job_skills: List of skills required by job
        
    Returns:
        List of matching skills (prioritized)
    """
    if not job_skills:
        return user_skills
        
    matching = []
    non_matching = []
    
    for skill in user_skills:
        skill_lower = skill.lower()
        # Check if any job skill matches
        if any(skill_lower in js.lower() or js.lower() in skill_lower 
               for js in job_skills):
            matching.append(skill)
        else:
            non_matching.append(skill)
    
    # Return matching skills first, then others
    return matching + non_matching


def create_tailored_resume(user_details: dict, base_resume_data: dict,
                          job_description: str, extracted_skills: dict,
                          output_format: str = 'pdf', output_path: str = None) -> str:
    """
    Creates a tailored resume based on job description and extracted skills.
    
    Args:
        user_details: Dict with user info
        base_resume_data: Dict with 'summary', 'experience', 'projects', 'skills', 'certificates'
        job_description: The job description text
        extracted_skills: Dict with skill categories from AI extraction
        output_format: 'pdf' or 'docx'
        output_path: Optional output path
        
    Returns:
        Path to the created resume file
    """
    # Get all required skills from extraction
    required_skills = []
    if extracted_skills:
        required_skills.extend(extracted_skills.get('required_skills', []))
        required_skills.extend(extracted_skills.get('tech_stack', []))
        required_skills.extend(extracted_skills.get('technical_skills', []))
    
    # Tailor the summary
    tailored_summary = tailor_resume_summary(
        base_resume_data.get('summary', ''),
        job_description,
        required_skills
    )
    
    # Reorder skills to highlight matches
    reordered_skills = highlight_matching_skills(
        base_resume_data.get('skills', []),
        required_skills
    )
    
    # Generate timestamp for unique filename
    if output_path is None:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = f"all resumes/temp/resume_{timestamp}.{output_format}"
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # Generate the resume
    if output_format.lower() == 'docx':
        return create_resume_docx_file(
            user_details,
            tailored_summary,
            base_resume_data.get('experience', []),
            base_resume_data.get('projects', []),
            reordered_skills,
            base_resume_data.get('certificates', []),
            output_path
        )
    else:
        return create_resume_pdf_file(
            user_details,
            tailored_summary,
            base_resume_data.get('experience', []),
            base_resume_data.get('projects', []),
            reordered_skills,
            base_resume_data.get('certificates', []),
            output_path
        )


def create_resume_docx_file(user_details, summary, experience, projects, skills, certificates, output_path):
    """Creates a DOCX resume file at the specified path."""
    doc = docx.Document()
    
    # Add user details
    doc.add_heading(user_details.get('name', 'Name'), 0)
    contact = f"{user_details.get('email', '')} | {user_details.get('phone_number', '')} | {user_details.get('address', '')}"
    doc.add_paragraph(contact)

    # Add summary
    doc.add_heading('Summary', 1)
    doc.add_paragraph(summary)

    # Add experience
    doc.add_heading('Experience', 1)
    for exp in experience:
        doc.add_heading(exp.get('company', ''), 2)
        doc.add_paragraph(f"{exp.get('role', '')} | {exp.get('dates', '')}")
        doc.add_paragraph(exp.get('achievements', ''))

    # Add projects
    doc.add_heading('Projects', 1)
    for project in projects:
        doc.add_heading(project.get('name', ''), 2)
        doc.add_paragraph(f"{project.get('description', '')} | {project.get('technologies', '')}")

    # Add skills
    doc.add_heading('Skills', 1)
    doc.add_paragraph(', '.join(skills))

    # Add certificates
    doc.add_heading('Certificates', 1)
    for cert in certificates:
        doc.add_heading(cert.get('name', ''), 2)
        doc.add_paragraph(cert.get('description', ''))

    doc.save(output_path)
    return output_path


def create_resume_pdf_file(user_details, summary, experience, projects, skills, certificates, output_path):
    """Creates a PDF resume file at the specified path."""
    pdf = FPDF()
    pdf.add_page()
    
    # Header with name
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, user_details.get('name', 'Name'), 0, 1, 'C')
    
    # Contact info
    pdf.set_font('Arial', '', 10)
    contact = f"{user_details.get('email', '')} | {user_details.get('phone_number', '')} | {user_details.get('address', '')}"
    pdf.cell(0, 8, contact, 0, 1, 'C')
    pdf.ln(5)

    # Summary
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'Summary', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(0, 5, summary)
    pdf.ln(5)

    # Experience
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'Experience', 0, 1, 'L')
    for exp in experience:
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 6, exp.get('company', ''), 0, 1, 'L')
        pdf.set_font('Arial', 'I', 10)
        pdf.cell(0, 6, f"{exp.get('role', '')} | {exp.get('dates', '')}", 0, 1, 'L')
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, exp.get('achievements', ''))
        pdf.ln(3)

    # Projects
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'Projects', 0, 1, 'L')
    for project in projects:
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 6, project.get('name', ''), 0, 1, 'L')
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, f"{project.get('description', '')} | {project.get('technologies', '')}")
        pdf.ln(3)

    # Skills
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'Skills', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    pdf.multi_cell(0, 5, ', '.join(skills))
    pdf.ln(5)

    # Certificates
    pdf.set_font('Arial', 'B', 12)
    pdf.cell(0, 8, 'Certificates', 0, 1, 'L')
    for cert in certificates:
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 6, cert.get('name', ''), 0, 1, 'L')
        pdf.set_font('Arial', '', 10)
        pdf.multi_cell(0, 5, cert.get('description', ''))
        pdf.ln(2)

    pdf.output(output_path, 'F')
    return output_path


# Legacy function for backward compatibility
def create_resume_docx(user_details, summary, experience, projects, skills, certificates):
    # Create a docx file
    doc = docx.Document()

    # Add user details
    doc.add_heading(user_details['name'], 0)
    doc.add_paragraph(user_details['email'] + ' | ' + user_details['phone_number'] + ' | ' + user_details['address'])

    # Add summary
    doc.add_heading('Summary', 1)
    doc.add_paragraph(summary)

    # Add experience
    doc.add_heading('Experience', 1)
    for experience_item in experience:
        doc.add_heading(experience_item['company'], 2)
        doc.add_paragraph(experience_item['role'] + ' | ' + experience_item['dates'])
        doc.add_paragraph(experience_item['achievements'])

    # Add projects
    doc.add_heading('Projects', 1)
    for project in projects:
        doc.add_heading(project['name'], 2)
        doc.add_paragraph(project['description'] + ' | ' + project['technologies'])

    # Add skills
    doc.add_heading('Skills', 1)
    doc.add_paragraph(', '.join(skills))

    # Add certificates
    doc.add_heading('Certificates', 1)
    for certificate in certificates:
        doc.add_heading(certificate['name'], 2)
        doc.add_paragraph(certificate['description'])

    # Save docx file
    doc.save('resume.docx')

    # Create a pdf file
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 12)

    # Add user details
    pdf.cell(0, 10, user_details['name'], 0, 1, 'C')
    pdf.cell(0, 10, user_details['email'] + ' | ' + user_details['phone_number'] + ' | ' + user_details['address'], 0, 1, 'L')

    # Add summary
    pdf.cell(0, 10, 'Summary', 0, 1, 'L')
    pdf.multi_cell(0, 10, summary)

    # Add experience
    pdf.cell(0, 10, 'Experience', 0, 1, 'L')
    for experience_item in experience:
        pdf.cell(0, 10, experience_item['company'], 0, 1, 'L')
        pdf.cell(0, 10, experience_item['role'] + ' | ' + experience_item['dates'], 0, 1, 'L')
        pdf.multi_cell(0, 10, experience_item['achievements'])

    # Add projects
    pdf.cell(0, 10, 'Projects', 0, 1, 'L')
    for project in projects:
        pdf.cell(0, 10, project['name'], 0, 1, 'L')
        pdf.multi_cell(0, 10, project['description'] + ' | ' + project['technologies'])

    # Add skills
    pdf.cell(0, 10, 'Skills', 0, 1, 'L')
    pdf.multi_cell(0, 10, ', '.join(skills))

    # Add certificates
    pdf.cell(0, 10, 'Certificates', 0, 1, 'L')
    for certificate in certificates:
        pdf.cell(0, 10, certificate['name'], 0, 1, 'L')
        pdf.multi_cell(0, 10, certificate['description'])

    # Save pdf file
    pdf.output('resume.pdf', 'F')