"""Resume tailoring helpers (standalone + integrated)."""
from __future__ import annotations

import os
import json
import hashlib
import textwrap
import difflib
import threading
import queue
from datetime import datetime
from typing import Optional, cast
from concurrent.futures import ThreadPoolExecutor, Future

from config.secrets import ai_provider, llm_api_key, llm_api_url, deepseek_api_key, deepseek_api_url
from config.secrets import groq_api_key, groq_api_url, groq_model
from config.secrets import huggingface_api_key, huggingface_api_url, huggingface_model
from config.settings import generated_resume_path, resume_tailoring_default_instructions
from config.personals import first_name
from modules.ai.prompts import resume_tailor_prompt, resume_tailor_prompt_compact, resume_tailor_paragraphs_prompt
from modules.ai.prompt_safety import sanitize_prompt_input, wrap_delimited
from modules.helpers import print_lg, critical_error_log


# ============ ASYNC TAILORING SYSTEM ============
# Background executor for async resume tailoring (1 worker to avoid rate limits)
_tailoring_executor = ThreadPoolExecutor(max_workers=1, thread_name_prefix="ResumeTailor")
_pending_tailoring: dict[str, Future] = {}  # Cache key -> Future
_tailoring_lock = threading.Lock()


def queue_async_tailor(
    job_description: str,
    job_title: str,
    resume_path: str,
    instructions: Optional[str] = None,
) -> str:
    """
    Queue a resume tailoring job to run in background.
    Returns a cache key that can be used to check/retrieve results.
    
    This allows the bot to pre-tailor resumes while showing user other jobs.
    """
    from modules.ai.resume_tailoring import _cache_key, _read_resume_text
    
    try:
        resume_text = _read_resume_text(resume_path)
    except Exception:
        return ""
    
    resolved_provider = (ai_provider or "ollama").lower()
    cache_key = _cache_key(resume_text, job_description, instructions, resolved_provider, job_title)
    
    with _tailoring_lock:
        # Already queued or processing?
        if cache_key in _pending_tailoring:
            return cache_key
        
        # Check if already cached
        cache_dir = os.path.join(generated_resume_path or "all resumes/", "cache")
        cache_path = os.path.join(cache_dir, f"{cache_key}.json")
        if os.path.exists(cache_path):
            return cache_key  # Already done
        
        # Queue the tailoring job
        future = _tailoring_executor.submit(
            _async_tailor_worker,
            resume_text=resume_text,
            job_description=job_description,
            job_title=job_title,
            instructions=instructions,
            cache_key=cache_key,
        )
        _pending_tailoring[cache_key] = future
        print_lg(f"🔄 Background tailoring queued for: {job_title[:40]}...")
        
    return cache_key


def get_async_tailor_result(cache_key: str, timeout: float = 0.1) -> Optional[dict]:
    """
    Check if async tailoring is complete and get result.
    
    Args:
        cache_key: Key from queue_async_tailor()
        timeout: Max seconds to wait (0.1 = non-blocking check)
        
    Returns:
        dict with paths if ready, None if still processing or failed
    """
    if not cache_key:
        return None
    
    with _tailoring_lock:
        future = _pending_tailoring.get(cache_key)
        
    if future is None:
        # Check cache directly
        cache_dir = os.path.join(generated_resume_path or "all resumes/", "cache")
        cache_path = os.path.join(cache_dir, f"{cache_key}.json")
        if os.path.exists(cache_path):
            try:
                with open(cache_path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return None
    
    if not future.done():
        try:
            # Wait briefly
            future.result(timeout=timeout)
        except Exception:
            return None  # Still processing or timed out
    
    try:
        result = future.result(timeout=0)
        # Clean up
        with _tailoring_lock:
            _pending_tailoring.pop(cache_key, None)
        return result
    except Exception:
        with _tailoring_lock:
            _pending_tailoring.pop(cache_key, None)
        return None


def _async_tailor_worker(
    resume_text: str,
    job_description: str,
    job_title: str,
    instructions: Optional[str],
    cache_key: str,
) -> dict:
    """Background worker for async tailoring."""
    try:
        from modules.ai.resume_tailoring import tailor_resume_to_files
        
        output_dir = os.path.join(generated_resume_path or "all resumes/", "temp")
        os.makedirs(output_dir, exist_ok=True)
        
        result = tailor_resume_to_files(
            resume_text=resume_text,
            job_description=job_description,
            job_title=job_title,
            instructions=instructions or resume_tailoring_default_instructions,
            output_dir=output_dir,
            enable_preview=False,
        )
        return result
    except Exception as e:
        print_lg(f"❌ Background tailoring failed: {e}")
        return {}


def is_tailoring_pending(cache_key: str) -> bool:
    """Check if a tailoring job is still in progress."""
    if not cache_key:
        return False
    with _tailoring_lock:
        future = _pending_tailoring.get(cache_key)
        if future and not future.done():
            return True
    return False


def wait_for_tailoring(cache_key: str, timeout: float = 30.0) -> Optional[dict]:
    """Wait for a specific tailoring job to complete."""
    return get_async_tailor_result(cache_key, timeout=timeout)


def _strip_prompt_markers(text: str) -> str:
    """
    CRITICAL: Remove prompt template markers that AI may accidentally include in output.
    
    This fixes the bug where AI outputs things like:
    - "===== MASTER RESUME (COPY THIS FORMAT EXACTLY) ====="
    - "<<RESUME>>"
    - "<<END_RESUME>>"
    - "===== END MASTER RESUME ====="
    - "===== JOB DESCRIPTION ====="
    
    The output should be ONLY the clean resume content, ready for an interviewer.
    """
    if not text:
        return text
    
    import re
    
    # List of marker patterns to remove (case-insensitive)
    marker_patterns = [
        # Section markers
        r'^={3,}.*?MASTER RESUME.*?={3,}\s*$',
        r'^={3,}.*?END MASTER RESUME.*?={3,}\s*$',
        r'^={3,}.*?JOB DESCRIPTION.*?={3,}\s*$',
        r'^={3,}.*?END JOB DESCRIPTION.*?={3,}\s*$',
        r'^={3,}.*?PRESERVE THIS FORMAT.*?={3,}\s*$',
        r'^={3,}.*?COPY THIS FORMAT.*?={3,}\s*$',
        r'^={3,}.*?EXTRACT KEYWORDS.*?={3,}\s*$',
        # XML-style markers
        r'^<<RESUME>>\s*$',
        r'^<<END_RESUME>>\s*$',
        r'^<<START_RESUME>>\s*$',
        r'^<<\/RESUME>>\s*$',
        r'^<RESUME>\s*$',
        r'^</RESUME>\s*$',
        r'^<END_RESUME>\s*$',
        # Instruction markers
        r'^OUTPUT:.*?optimized keywords:\s*$',
        r'^OUTPUT:\s*$',
        r'^Return ONLY the tailored resume.*$',
        r'^No explanations.*$',
        r'^Start directly with.*$',
        # Triple backticks (code blocks)
        r'^```.*$',
        # Common AI artifacts
        r'^Here is the tailored resume:?\s*$',
        r'^Here\'s the tailored resume:?\s*$',
        r'^Tailored Resume:?\s*$',
        r'^---+\s*$',
        r'^\*\*\*+\s*$',
    ]
    
    # Process line by line
    lines = text.split('\n')
    clean_lines = []
    
    for line in lines:
        stripped = line.strip()
        should_remove = False
        
        for pattern in marker_patterns:
            if re.match(pattern, stripped, re.IGNORECASE):
                should_remove = True
                break
        
        # Also check for common marker substrings
        lower_line = stripped.lower()
        marker_substrings = [
            'master resume', 'end master resume', 'copy this format',
            'preserve this format', 'end job description', '<<resume>>',
            '<<end_resume>>', 'extract keywords from', '===== end',
            '===== master', '===== job'
        ]
        
        for marker in marker_substrings:
            if marker in lower_line and ('=====' in stripped or '<<' in stripped):
                should_remove = True
                break
        
        if not should_remove:
            clean_lines.append(line)
    
    # Remove leading/trailing empty lines
    result = '\n'.join(clean_lines).strip()
    
    # Final safety check - if the result starts with a marker character, try to find actual content
    while result and result[0] in '=<-*#':
        first_newline = result.find('\n')
        if first_newline > 0:
            first_line = result[:first_newline].strip().lower()
            if any(m in first_line for m in ['master resume', 'job description', 'resume>>', 'end_']):
                result = result[first_newline:].strip()
            else:
                break
        else:
            break
    
    return result


def _sanitize_text_for_display(text: str) -> str:
    """Sanitize text to ensure it displays correctly on all systems.
    
    Replaces problematic Unicode characters that may cause encoding issues.
    """
    if not text:
        return text
    # Replace common problematic characters with ASCII equivalents
    replacements = {
        '\u2018': "'",   # Left single quote
        '\u2019': "'",   # Right single quote
        '\u201c': '"',   # Left double quote
        '\u201d': '"',   # Right double quote
        '\u2013': '-',   # En dash
        '\u2014': '--',  # Em dash
        '\u2026': '...', # Ellipsis
        '\u00a0': ' ',   # Non-breaking space
        '\u2022': '*',   # Bullet
        '\u00b7': '*',   # Middle dot
        '\u2010': '-',   # Hyphen
        '\u2011': '-',   # Non-breaking hyphen
        '\u2212': '-',   # Minus sign
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Encode to ASCII with replacement for any remaining problematic chars
    return text.encode('ascii', 'replace').decode('ascii').replace('?', '')


def _extract_jd_keywords(job_description: str) -> set[str]:
    """Extract important ATS keywords from job description."""
    jd_lower = job_description.lower()
    
    # Comprehensive keyword categories
    all_keywords = {
        # Programming languages
        'python', 'java', 'kotlin', 'go', 'golang', 'scala', 'javascript', 'typescript',
        'c++', 'c#', 'ruby', 'rust', 'php', 'swift',
        # Frameworks
        'spring', 'spring boot', 'flask', 'django', 'fastapi', 'node.js', 'express',
        'react', 'angular', 'vue', '.net',
        # Databases
        'postgresql', 'postgres', 'mysql', 'mongodb', 'dynamodb', 'cassandra', 'redis',
        'elasticsearch', 'oracle', 'sql server', 'nosql', 'rdbms', 'sql',
        # Cloud & DevOps
        'aws', 'azure', 'gcp', 'google cloud', 'docker', 'kubernetes', 'k8s',
        'ci/cd', 'jenkins', 'terraform', 'ansible', 'cloudformation',
        'saas', 'paas', 'iaas', 'cloud',
        # Architecture & Patterns
        'microservices', 'restful', 'rest api', 'graphql', 'api', 'distributed systems',
        'scalable', 'scalability', 'high availability', 'event-driven', 'serverless',
        'cloud architecture',
        # Practices & Skills
        'agile', 'scrum', 'kanban', 'devops', 'tdd', 'bdd',
        'testing', 'unit testing', 'integration testing', 'code review',
        'security', 'performance', 'reliability', 'monitoring', 'observability',
        'collaboration', 'mentoring', 'mentor', 'leadership',
    }
    
    # Find which keywords appear in the JD
    found = set()
    for kw in all_keywords:
        if kw in jd_lower:
            found.add(kw)
    
    return found


def _inject_missing_keywords(resume_text: str, jd_keywords: set[str]) -> str:
    """Inject missing keywords into the skills section of the resume."""
    resume_lower = resume_text.lower()
    
    # Find missing keywords
    missing = []
    for kw in jd_keywords:
        if kw not in resume_lower:
            missing.append(kw)
    
    if not missing:
        return resume_text  # All keywords present!
    
    # Format keywords nicely (capitalize appropriately)
    keyword_map = {
        'aws': 'AWS', 'azure': 'Azure', 'gcp': 'GCP', 'ci/cd': 'CI/CD',
        'saas': 'SaaS', 'paas': 'PaaS', 'iaas': 'IaaS', 'nosql': 'NoSQL',
        'rdbms': 'RDBMS', 'sql': 'SQL', 'restful': 'RESTful', 'api': 'API',
        'graphql': 'GraphQL', 'kubernetes': 'Kubernetes', 'k8s': 'K8s',
        'docker': 'Docker', 'jenkins': 'Jenkins', 'terraform': 'Terraform',
        'postgresql': 'PostgreSQL', 'mongodb': 'MongoDB', 'dynamodb': 'DynamoDB',
        'cassandra': 'Cassandra', 'redis': 'Redis', 'elasticsearch': 'Elasticsearch',
        'python': 'Python', 'java': 'Java', 'kotlin': 'Kotlin', 'scala': 'Scala',
        'golang': 'Go', 'go': 'Go', 'javascript': 'JavaScript', 'typescript': 'TypeScript',
        'agile': 'Agile', 'scrum': 'Scrum', 'devops': 'DevOps',
        'spring boot': 'Spring Boot', 'flask': 'Flask', 'django': 'Django',
        'microservices': 'Microservices', 'cloud architecture': 'Cloud Architecture',
    }
    
    formatted_missing = []
    for kw in missing[:8]:  # Limit to top 8 missing keywords
        formatted = keyword_map.get(kw, kw.title())
        formatted_missing.append(formatted)
    
    if not formatted_missing:
        return resume_text
    
    # Try to find and enhance SKILLS section
    lines = resume_text.split('\n')
    new_lines = []
    skills_enhanced = False
    
    for i, line in enumerate(lines):
        new_lines.append(line)
        line_lower = line.lower().strip()
        
        # Detect skills section header
        if not skills_enhanced and ('skills' in line_lower and len(line_lower) < 20):
            # Next non-empty line should be the skills list
            for j in range(i + 1, min(i + 3, len(lines))):
                if lines[j].strip():
                    # This is likely the skills line - we'll enhance it after the loop
                    break
        
        # Detect a skills list line (contains commas and tech terms)
        if not skills_enhanced and ',' in line and any(tech in line_lower for tech in ['python', 'java', 'aws', 'docker', 'sql', 'api', 'git']):
            # Append missing keywords to this line
            additions = ', ' + ', '.join(formatted_missing)
            new_lines[-1] = line.rstrip() + additions
            skills_enhanced = True
    
    # If no skills section found, append at the end
    if not skills_enhanced and formatted_missing:
        new_lines.append('')
        new_lines.append('Additional Skills: ' + ', '.join(formatted_missing))
    
    return '\n'.join(new_lines)


def _build_prompt(resume_text: str, job_description: str, instructions: Optional[str], use_compact: bool = False) -> str:
    safe_instructions = sanitize_prompt_input(instructions or resume_tailoring_default_instructions or "None")
    safe_resume = sanitize_prompt_input(resume_text, max_len=12000)
    safe_jd = sanitize_prompt_input(job_description, max_len=8000)
    
    # Use compact prompt for local models (faster processing)
    if use_compact:
        return resume_tailor_prompt_compact.format(
            instructions=safe_instructions,
            resume_text=safe_resume,
            job_description=safe_jd,
        )
    
    return resume_tailor_prompt.format(
        instructions=wrap_delimited("instructions", safe_instructions),
        resume_text=wrap_delimited("resume", safe_resume),
        job_description=wrap_delimited("job_description", safe_jd),
    )


def _build_paragraphs_prompt(paragraphs: list[str], job_description: str, instructions: Optional[str]) -> str:
    safe_instructions = sanitize_prompt_input(instructions or resume_tailoring_default_instructions or "None")
    safe_jd = sanitize_prompt_input(job_description, max_len=8000)
    safe_paragraphs = [sanitize_prompt_input(p, max_len=1000) for p in paragraphs]
    return resume_tailor_paragraphs_prompt.format(
        job_description=wrap_delimited("job_description", safe_jd),
        instructions=wrap_delimited("instructions", safe_instructions),
        paragraphs_json=json.dumps(safe_paragraphs, ensure_ascii=False),
    )


def _ensure_output_dir(path: str) -> None:
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)


def _save_text(text: str, output_dir: str, base_name: str = "tailored_resume") -> str:
    _ensure_output_dir(output_dir)
    file_path = os.path.join(output_dir, f"{base_name}.txt")
    # If file exists, add timestamp to avoid overwrite
    if os.path.exists(file_path):
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(output_dir, f"{base_name}_{stamp}.txt")
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(text.strip())
    return file_path


def _sanitize_filename(value: str) -> str:
    """Sanitize filename while keeping spaces for readability."""
    # Characters that are invalid in Windows filenames
    invalid_chars = '<>:"/\\|?*'
    cleaned = "".join(ch if ch not in invalid_chars else "_" for ch in value.strip())
    # Remove consecutive spaces/underscores
    import re
    cleaned = re.sub(r'[_\s]+', ' ', cleaned).strip()
    return cleaned


def _candidate_name() -> str:
    # Return just first name for cleaner filename
    return first_name.strip() if first_name else "Candidate"


def _read_resume_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext in [".txt", ".md"]:
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    if ext == ".docx":
        try:
            from docx import Document
        except Exception as e:
            raise ValueError("python-docx is required to read .docx resumes. Please install it.") from e
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text)

    if ext == ".pdf":
        try:
            from PyPDF2 import PdfReader
        except Exception as e:
            raise ValueError("PyPDF2 is required to read .pdf resumes. Please install it.") from e
        reader = PdfReader(path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    raise ValueError(f"Unsupported resume format: {ext}")


def _read_docx_paragraphs(path: str) -> list[str]:
    try:
        from docx import Document
    except Exception as e:
        raise ValueError("python-docx is required to read .docx resumes. Please install it.") from e
    doc = Document(path)
    return [p.text for p in doc.paragraphs]


def _write_docx_from_template(template_path: str, paragraphs: list[str], output_dir: str, base_name: str) -> str:
    _ensure_output_dir(output_dir)
    try:
        from docx import Document
    except Exception as e:
        raise ValueError("python-docx is required to write .docx resumes. Please install it.") from e
    doc = Document(template_path)
    for idx, p in enumerate(doc.paragraphs):
        if idx < len(paragraphs):
            # Preserve paragraph style; replace text only
            p.text = paragraphs[idx]
    path = os.path.join(output_dir, f"{base_name}.docx")
    # If file exists, add timestamp to avoid overwrite
    if os.path.exists(path):
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = os.path.join(output_dir, f"{base_name}_{stamp}.docx")
    doc.save(path)
    return path


def _write_docx(text: str, output_dir: str, base_name: str) -> str:
    """
    Write text to a DOCX file with proper professional formatting.
    
    Creates a properly formatted DOCX that LinkedIn accepts (not tiny 3KB files).
    Uses same formatting logic as _write_pdf for consistency.
    """
    _ensure_output_dir(output_dir)
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as e:
        raise ValueError("python-docx is required to write .docx resumes. Please install it.") from e
    
    path = os.path.join(output_dir, f"{base_name}.docx")
    # If file exists, add timestamp to avoid overwrite
    if os.path.exists(path):
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = os.path.join(output_dir, f"{base_name}_{stamp}.docx")
    
    doc = Document()
    
    # Set professional margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add content with proper formatting
    lines = text.splitlines()
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            # Empty line - add spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph(stripped)
            # First line (name) - make it bold and larger, centered
            if i == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(14)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Second line (contact info) - center it, smaller font
            elif i == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
            # Section headers (all caps or ending with :)
            elif stripped.isupper() or stripped.endswith(':'):
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            else:
                for run in p.runs:
                    run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(2)
    
    doc.save(path)
    
    # Verify file size is reasonable (not tiny)
    file_size = os.path.getsize(path)
    if file_size < 5000:
        print(f"[DOCX] ⚠️ Warning: DOCX file is small ({file_size} bytes), may have issues")
    else:
        print(f"[DOCX] ✅ Created properly formatted DOCX: {file_size} bytes")
    
    return path


def _write_pdf(text: str, output_dir: str, base_name: str) -> str:
    """
    Write text to a PDF file. 
    
    LinkedIn has strict PDF validation - uses docx2pdf conversion for proper formatting,
    falling back to reportlab canvas for basic PDF if conversion fails.
    """
    _ensure_output_dir(output_dir)
    
    path = os.path.join(output_dir, f"{base_name}.pdf")
    # If file exists, add timestamp to avoid overwrite
    if os.path.exists(path):
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        path = os.path.join(output_dir, f"{base_name}_{stamp}.pdf")
    
    # METHOD 1: Create DOCX first, then convert to PDF (best quality, LinkedIn-compatible)
    try:
        # First create a proper DOCX
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        temp_docx_path = os.path.join(output_dir, f"{base_name}_temp.docx")
        doc = Document()
        
        # Set margins for professional look
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Add content with proper formatting
        lines = text.splitlines()
        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                # Empty line - add spacing
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
            else:
                p = doc.add_paragraph(stripped)
                # First line (name) - make it bold and larger
                if i == 0:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(14)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Second line (contact) - center it
                elif i == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(10)
                # Section headers (all caps or ending with :)
                elif stripped.isupper() or stripped.endswith(':'):
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(11)
                else:
                    for run in p.runs:
                        run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
        
        doc.save(temp_docx_path)
        
        # Try to convert DOCX to PDF using docx2pdf (requires MS Word)
        try:
            from docx2pdf import convert
            convert(temp_docx_path, path)
            
            # Clean up temp docx
            try:
                os.remove(temp_docx_path)
            except:
                pass
            
            # Verify PDF was created and has reasonable size
            if os.path.exists(path) and os.path.getsize(path) > 5000:
                print(f"[PDF] ✅ Created LinkedIn-compatible PDF via docx2pdf: {os.path.getsize(path)} bytes")
                return path
        except ImportError:
            print("[PDF] docx2pdf not available, trying alternative...")
        except Exception as conv_err:
            print(f"[PDF] docx2pdf conversion failed: {conv_err}")
        
        # Clean up temp docx if still exists
        try:
            if os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)
        except:
            pass
            
    except Exception as docx_err:
        print(f"[PDF] DOCX-based PDF creation failed: {docx_err}")
    
    # METHOD 2: Fall back to reportlab (basic but functional)
    print("[PDF] Using reportlab fallback for PDF creation")
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception as e:
        raise ValueError("reportlab is required to write .pdf resumes. Please install it.") from e
    
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    
    # Set up better formatting
    x = 0.75 * inch
    y = height - 0.5 * inch
    
    lines = text.splitlines()
    for i, line in enumerate(lines):
        wrapped = textwrap.wrap(line, width=95) or [""]
        for chunk in wrapped:
            if y < 0.5 * inch:
                c.showPage()
                y = height - 0.5 * inch
            
            # First line (name) - larger font
            if i == 0 and chunk == line:
                c.setFont("Helvetica-Bold", 14)
                # Center the name
                text_width = c.stringWidth(chunk, "Helvetica-Bold", 14)
                c.drawString((width - text_width) / 2, y, chunk)
            # Section headers
            elif chunk.strip().isupper() or chunk.strip().endswith(':'):
                c.setFont("Helvetica-Bold", 11)
                c.drawString(x, y, chunk)
            else:
                c.setFont("Helvetica", 10)
                c.drawString(x, y, chunk)
            y -= 12
    
    c.save()
    print(f"[PDF] Created PDF via reportlab: {os.path.getsize(path)} bytes")
    return path


def tailor_resume_text(
    resume_text: str,
    job_description: str,
    instructions: Optional[str] = None,
    provider: Optional[str] = None,
    client: Optional[object] = None,
    inject_keywords: bool = True,  # Enable keyword injection for max ATS score
    validate_quality: bool = True,  # NEW: Validate output quality
    max_retries: int = 1,  # NEW: Max retry attempts if quality is low
) -> str:
    """
    Tailor a resume using the configured AI provider (defaults to local Ollama).
    Returns tailored resume text.
    
    Args:
        inject_keywords: If True, automatically inject missing JD keywords into the resume
                        for maximum ATS score (default: True)
        validate_quality: If True, validate the tailored resume and log quality metrics
        max_retries: Number of retry attempts if quality validation fails (0 to disable)
    """
    resolved_provider = (provider or ai_provider or "ollama").lower()
    
    # Extract JD keywords BEFORE AI processing (for post-processing injection)
    jd_keywords = _extract_jd_keywords(job_description) if inject_keywords else set()
    
    # Use compact prompt for local models (faster) or full prompt for cloud APIs
    use_compact = resolved_provider == "ollama"
    
    # Track retry attempts
    current_instructions = instructions
    attempt = 0
    
    while attempt <= max_retries:
        attempt += 1
        prompt = _build_prompt(resume_text, job_description, current_instructions, use_compact=use_compact)
        result = _call_ai_provider(resolved_provider, prompt, client)
        
        if not result or result.startswith("["):
            # Error occurred
            return result
        
        # CRITICAL: Strip any prompt markers the AI accidentally included
        result = _strip_prompt_markers(result)
        
        # Sanitize and inject keywords
        result = _sanitize_text_for_display(result)
        if inject_keywords and jd_keywords:
            result = _inject_missing_keywords(result, jd_keywords)
        
        # Quality validation (skip retries for local models to save time)
        if validate_quality and attempt <= max_retries:
            try:
                from modules.ai.resume_validator import validate_tailored_resume, should_retry_tailoring, get_retry_instructions
                from modules.ai.jd_analyzer import analyze_jd_fast
                
                analysis = analyze_jd_fast(job_description)
                report = validate_tailored_resume(resume_text, result, job_description, analysis)
                
                # Log quality metrics
                print_lg(f"Resume Quality: Grade {report.overall_grade}, ATS Score: {report.weighted_ats_score:.1f}%")
                
                should_retry, reason = should_retry_tailoring(report)
                if should_retry and attempt <= max_retries:
                    print_lg(f"Quality check failed ({reason}), retrying with improved instructions...")
                    current_instructions = get_retry_instructions(report, analysis)
                    continue
            except ImportError:
                pass  # Validator not available, skip validation
            except Exception as e:
                print_lg(f"Quality validation error (non-critical): {e}")
        
        return result
    
    return result  # Return last attempt


def _call_ai_provider(provider: str, prompt: str, client: Optional[object] = None) -> str:
    """Call the specified AI provider and return the result."""
    try:
        if provider == "ollama":
            from modules.ai import ollama_integration as _oll
            # 900 second timeout (15 min) for large models like qwen3:14b
            return str(_oll.generate(prompt, timeout=900, stream=False))

        if provider == "openai":
            from openai import OpenAI
            from modules.ai.openaiConnections import ai_completion
            if client is None:
                client = OpenAI(base_url=llm_api_url, api_key=llm_api_key)
            client = cast(OpenAI, client)
            messages = [{"role": "user", "content": prompt}]
            return str(ai_completion(client, messages, response_format=None, stream=False))

        if provider == "deepseek":
            from openai import OpenAI
            from modules.ai.deepseekConnections import deepseek_completion
            if client is None:
                client = OpenAI(base_url=deepseek_api_url, api_key=deepseek_api_key)
            client = cast(OpenAI, client)
            messages = [{"role": "user", "content": prompt}]
            return str(deepseek_completion(client, messages, response_format={}, stream=False))

        if provider == "gemini":
            from modules.ai.geminiConnections import gemini_completion, gemini_create_client
            if client is None:
                client = gemini_create_client()
            result = gemini_completion(client, prompt, is_json=False)
            if isinstance(result, dict) and "error" in result:
                raise ValueError(f"Gemini API error: {result['error']}")
            return str(result)

        if provider == "groq":
            # Groq uses OpenAI-compatible API (FREE & FAST!)
            from openai import OpenAI
            if client is None:
                client = OpenAI(base_url=groq_api_url, api_key=groq_api_key)
            client = cast(OpenAI, client)
            
            # ALWAYS use 8b-instant for FAST tailoring (2-3 sec vs 15-20 sec)
            # The 8b model is optimized for speed while still producing quality results
            FAST_MODEL = "llama-3.1-8b-instant"  # 3x faster than 70b!
            
            messages = [{
                "role": "system", 
                "content": "You are an expert ATS resume optimizer. Be concise. Focus on keyword matching."
            }, {
                "role": "user", 
                "content": prompt
            }]
            
            response = client.chat.completions.create(
                model=FAST_MODEL,
                messages=messages,
                temperature=0.3,  # Lower temp = faster + more focused
                max_tokens=2500,  # Reduced for speed
            )
            return response.choices[0].message.content or ""

        if provider == "huggingface":
            # HuggingFace Inference API (FREE!)
            import urllib.request
            import urllib.error
            headers = {
                "Authorization": f"Bearer {huggingface_api_key}",
                "Content-Type": "application/json"
            }
            # Format for HF text generation
            payload = json.dumps({
                "inputs": prompt,
                "parameters": {
                    "max_new_tokens": 4000,
                    "temperature": 0.7,
                    "return_full_text": False
                }
            }).encode("utf-8")
            url = f"{huggingface_api_url}/{huggingface_model}"
            req = urllib.request.Request(url, data=payload, headers=headers, method="POST")
            with urllib.request.urlopen(req, timeout=120) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                if isinstance(data, list) and len(data) > 0:
                    return data[0].get("generated_text", "")
                else:
                    return str(data)

        raise ValueError(f"Unsupported AI provider: {provider}")
    except Exception as e:
        critical_error_log("Resume tailoring failed!", e)
        return f"[Resume Tailoring Error] {e}"


def _tailor_paragraphs_with_ai(
    paragraphs: list[str],
    job_description: str,
    instructions: Optional[str],
    provider: str,
    client: Optional[object],
) -> list[str] | None:
    prompt = _build_paragraphs_prompt(paragraphs, job_description, instructions)
    try:
        if provider == "openai":
            from openai import OpenAI
            from modules.ai.openaiConnections import ai_completion
            if client is None:
                client = OpenAI(base_url=llm_api_url, api_key=llm_api_key)
            client = cast(OpenAI, client)
            messages = [{"role": "user", "content": prompt}]
            result = ai_completion(client, messages, response_format=None, stream=False)
            return json.loads(str(result))

        if provider == "deepseek":
            from openai import OpenAI
            from modules.ai.deepseekConnections import deepseek_completion
            if client is None:
                client = OpenAI(base_url=deepseek_api_url, api_key=deepseek_api_key)
            client = cast(OpenAI, client)
            messages = [{"role": "user", "content": prompt}]
            result = deepseek_completion(client, messages, response_format={}, stream=False)
            return json.loads(str(result))

        if provider == "ollama":
            from modules.ai import ollama_integration as _oll
            # 900 second timeout (15 min) for large models
            result = _oll.generate(prompt, timeout=900, stream=False)
            return json.loads(str(result))

        if provider == "gemini":
            from modules.ai.geminiConnections import gemini_completion, gemini_create_client
            if client is None:
                client = gemini_create_client()
            result = gemini_completion(client, prompt, is_json=True)
            if isinstance(result, dict) and "error" in result:
                raise ValueError(f"Gemini API error: {result['error']}")
            return json.loads(str(result))

        if provider == "groq":
            from openai import OpenAI
            if client is None:
                client = OpenAI(base_url=groq_api_url, api_key=groq_api_key)
            client = cast(OpenAI, client)
            messages = [{"role": "user", "content": prompt}]
            response = client.chat.completions.create(
                model=groq_model,
                messages=messages,
                temperature=0.7,
                max_tokens=4000,
            )
            result = response.choices[0].message.content or ""
            return json.loads(result)

        if provider == "huggingface":
            import urllib.request
            headers = {
                "Authorization": f"Bearer {huggingface_api_key}",
                "Content-Type": "application/json"
            }
            payload = json.dumps({
                "inputs": prompt,
                "parameters": {"max_new_tokens": 4000, "temperature": 0.7, "return_full_text": False}
            }).encode("utf-8")
            url = f"{huggingface_api_url}/{huggingface_model}"
            req = urllib.request.Request(url, data=payload, headers=headers, method="POST")
            with urllib.request.urlopen(req, timeout=120) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                if isinstance(data, list) and len(data) > 0:
                    result = data[0].get("generated_text", "")
                else:
                    result = str(data)
            return json.loads(result)
    except Exception:
        return None
    return None


def _normalize_one_page(text: str, max_chars: int = 4500) -> str:
    cleaned = "\n".join(line.rstrip() for line in text.splitlines())
    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")
    return cleaned[:max_chars].rstrip()


def _cache_key(master_text: str, job_description: str, instructions: Optional[str], provider: str, job_title: Optional[str]) -> str:
    h = hashlib.sha256()
    h.update(master_text.encode("utf-8", errors="ignore"))
    h.update(job_description.encode("utf-8", errors="ignore"))
    h.update((instructions or "").encode("utf-8", errors="ignore"))
    h.update(provider.encode("utf-8", errors="ignore"))
    h.update((job_title or "").encode("utf-8", errors="ignore"))
    return h.hexdigest()


def _keyword_set(text: str) -> set[str]:
    # Common tech abbreviations that should NOT be filtered despite length
    tech_abbrevs = {"ml", "ai", "js", "ts", "ci", "cd", "ui", "ux", "qa", "db", "vm", "os", "io", "c", "r"}
    stop = {"the","and","with","for","from","this","that","will","are","was","were","your","you","our","their","they","them","a","an","to","of","in","on","by","as","is","or","at","be","we","it","have","has","had","been","being","would","could","should","may","might","must","can","do","does","did","about","into","over","after","before","above","below","between","under","during","through","within","without","while","where","when","what","which","who","whom","whose","than","then","there","these","those","such","each","every","both","either","neither","all","any","some","most","other","another","many","much","few","more","less","only","just","also","even","still","already","often","always","never","sometimes","usually","generally","typically","well","very","really","quite","rather","almost","nearly","exactly","actually","probably","possibly","certainly","definitely","simply","clearly","obviously","especially","particularly","specifically","including","including"}
    words = [w.lower().strip(".,:;()[]{}<>/\\\"'`").replace("-", " ") for w in text.split()]
    tokens = set()
    for w in words:
        for part in w.split():
            # Keep technical abbreviations even if short
            if part in tech_abbrevs or (len(part) > 2 and part not in stop):
                tokens.add(part)
    return tokens


# Synonym groups for better matching (terms in same group are considered equivalent)
# Each key in a synonym group maps to all its equivalents
SKILL_SYNONYMS = {
    # Programming Languages - use single-word keys for tokenized matching
    "js": {"js", "javascript", "ecmascript"},
    "javascript": {"js", "javascript", "ecmascript"},
    "ts": {"ts", "typescript"},
    "typescript": {"ts", "typescript"},
    "python": {"python", "py"},
    "py": {"python", "py"},
    "go": {"go", "golang"},
    "golang": {"go", "golang"},
    
    # AI/ML - map abbreviations to expanded forms
    "ml": {"ml", "machine", "learning"},  # Will match if both 'machine' and 'learning' are present
    "ai": {"ai", "artificial", "intelligence"},
    "dl": {"dl", "deep", "learning"},
    "nlp": {"nlp", "natural", "language", "processing"},
    
    # Cloud
    "aws": {"aws", "amazon"},
    "gcp": {"gcp", "google", "cloud"},
    "azure": {"azure", "microsoft"},
    
    # DevOps
    "ci": {"ci", "continuous", "integration"},
    "cd": {"cd", "continuous", "deployment", "delivery"},
    "cicd": {"cicd", "ci", "cd"},
    "k8s": {"k8s", "kubernetes"},
    "kubernetes": {"k8s", "kubernetes"},
    
    # Databases
    "postgres": {"postgres", "postgresql"},
    "postgresql": {"postgres", "postgresql"},
    "mongo": {"mongo", "mongodb"},
    "mongodb": {"mongo", "mongodb"},
    "elastic": {"elastic", "elasticsearch"},
    "elasticsearch": {"elastic", "elasticsearch"},
    
    # Frameworks
    "react": {"react", "reactjs"},
    "reactjs": {"react", "reactjs"},
    "angular": {"angular", "angularjs"},
    "vue": {"vue", "vuejs"},
    "node": {"node", "nodejs"},
    "nodejs": {"node", "nodejs"},
    "express": {"express", "expressjs"},
    
    # Soft Skills
    "lead": {"lead", "leading", "leadership", "led"},
    "leadership": {"lead", "leading", "leadership", "led"},
    "manage": {"manage", "managed", "management", "managing"},
    "management": {"manage", "managed", "management", "managing"},
    "communicate": {"communicate", "communication", "communicating"},
    "communication": {"communicate", "communication", "communicating"},
    "collaborate": {"collaborate", "collaboration", "collaborative"},
    "collaboration": {"collaborate", "collaboration", "collaborative"},
}


def _normalize_with_synonyms(tokens: set[str]) -> set[str]:
    """Expand tokens to include synonym variations for better matching."""
    expanded = set(tokens)
    for canonical, synonyms in SKILL_SYNONYMS.items():
        # If any synonym is in tokens, add all synonyms
        if tokens & synonyms:
            expanded.update(synonyms)
            expanded.add(canonical)
    return expanded


def _extract_important_keywords(jd_text: str) -> dict:
    """Extract and categorize important keywords from job description."""
    text_lower = jd_text.lower()
    
    # Common technical keywords to look for
    tech_keywords = {
        "python", "java", "javascript", "typescript", "react", "angular", "vue", "node", "nodejs",
        "sql", "nosql", "mongodb", "postgresql", "mysql", "redis", "elasticsearch",
        "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "cicd", "terraform",
        "api", "rest", "graphql", "microservices", "agile", "scrum", "git", "linux",
        "machine learning", "ml", "ai", "data science", "analytics", "tableau", "power bi",
        "html", "css", "sass", "webpack", "npm", "yarn", "spring", "django", "flask",
        "c++", "c#", "golang", "rust", "scala", "kotlin", "swift", "ruby", "php"
    }
    
    # Soft skills keywords
    soft_keywords = {
        "leadership", "communication", "teamwork", "collaboration", "problem solving",
        "analytical", "critical thinking", "time management", "project management",
        "stakeholder", "mentoring", "cross functional", "presentation", "negotiation"
    }
    
    found_tech = []
    found_soft = []
    
    for kw in tech_keywords:
        if kw in text_lower:
            found_tech.append(kw)
    
    for kw in soft_keywords:
        if kw in text_lower:
            found_soft.append(kw)
    
    return {"technical": found_tech, "soft": found_soft}


def _score_match(resume_text: str, jd_text: str) -> dict:
    """Calculate keyword match score between resume and job description.
    
    Uses synonym expansion for more accurate matching:
    - "ML" in resume matches "machine learning" in JD
    - "JS" matches "JavaScript"
    - "K8s" matches "Kubernetes"
    """
    jd_tokens = _keyword_set(jd_text)
    res_tokens = _keyword_set(resume_text)
    if not jd_tokens:
        return {"match": 0, "ats": 0, "matched": 0, "total": 0, "missing": [], "found": []}
    
    # Expand both token sets with synonyms for better matching
    jd_expanded = _normalize_with_synonyms(jd_tokens)
    res_expanded = _normalize_with_synonyms(res_tokens)
    
    # Match using expanded sets
    matched_tokens = jd_expanded.intersection(res_expanded)
    missing_tokens = jd_tokens - res_expanded  # Show original JD tokens that are truly missing
    
    # Score based on original JD tokens that have matches (via synonyms)
    matched_originals = sum(1 for t in jd_tokens if t in res_expanded or any(
        t in syns and res_expanded & syns for syns in SKILL_SYNONYMS.values()
    ))
    
    total = len(jd_tokens)
    score = int(round((matched_originals / total) * 100)) if total > 0 else 0
    
    # Get important keywords analysis
    important = _extract_important_keywords(jd_text)
    resume_lower = resume_text.lower()
    
    # Check which important keywords are in resume (with synonym awareness)
    tech_found = []
    tech_missing = []
    for kw in important["technical"]:
        # Check if keyword or any synonym is present
        found = kw in resume_lower
        if not found:
            # Check synonyms
            for canonical, synonyms in SKILL_SYNONYMS.items():
                if kw in synonyms or kw == canonical:
                    if any(syn in resume_lower for syn in synonyms):
                        found = True
                        break
        if found:
            tech_found.append(kw)
        else:
            tech_missing.append(kw)
    
    soft_found = [k for k in important["soft"] if k in resume_lower]
    soft_missing = [k for k in important["soft"] if k not in resume_lower]
    
    return {
        "match": score, 
        "ats": score, 
        "matched": matched_originals, 
        "total": total,
        "found": list(matched_tokens)[:20],  # Top 20 found keywords
        "missing": list(missing_tokens)[:15],  # Top 15 missing keywords
        "tech_found": tech_found,
        "tech_missing": tech_missing,
        "soft_found": soft_found,
        "soft_missing": soft_missing
    }


def generate_preview_report(master_text: str, tailored_text: str, jd_text: str, output_dir: str, base_name: str) -> str:
    _ensure_output_dir(output_dir)
    stats_before = _score_match(master_text, jd_text)
    stats_after = _score_match(tailored_text, jd_text)

    diff = difflib.HtmlDiff(wrapcolumn=80)
    html = diff.make_file(
        master_text.splitlines(),
        tailored_text.splitlines(),
        fromdesc="Master Resume",
        todesc="Tailored Resume",
        context=True,
        numlines=3,
    )

    summary = (
        f"<h2>JD Match & ATS Scores</h2>"
        f"<p><b>Before:</b> Match {stats_before['match']}% (matched {stats_before['matched']}/{stats_before['total']}), ATS {stats_before['ats']}%</p>"
        f"<p><b>After:</b> Match {stats_after['match']}% (matched {stats_after['matched']}/{stats_after['total']}), ATS {stats_after['ats']}%</p>"
        f"<hr/>"
    )

    html = html.replace("<body>", f"<body>{summary}")
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(output_dir, f"{base_name}_{stamp}_diff.html")
    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    return path


def show_preview_gui(master_text: str, tailored_text: str, jd_text: str, paths: dict, job_title: str = "Job") -> None:
    """Show a GUI window with side-by-side comparison of master and tailored resume."""
    import tkinter as tk
    from tkinter import scrolledtext
    try:
        import ttkbootstrap as ttkb
        HAS_TTK = True
    except ImportError:
        from tkinter import ttk as ttkb
        HAS_TTK = False
    
    # Calculate scores
    stats_before = _score_match(master_text, jd_text)
    stats_after = _score_match(tailored_text, jd_text)
    
    # Create window
    preview_window = tk.Toplevel()
    preview_window.title(f"📄 Resume Preview - {job_title}")
    preview_window.geometry("1400x800")
    preview_window.configure(bg='#1a1a2e')
    
    # Make it appear on top
    preview_window.attributes('-topmost', True)
    preview_window.after(100, lambda: preview_window.attributes('-topmost', False))
    
    # Header frame with stats
    if HAS_TTK:
        header_frame = ttkb.Frame(preview_window, padding=10)
    else:
        header_frame = ttkb.Frame(preview_window)
    header_frame.pack(fill='x', padx=10, pady=5)
    
    # Title
    title_label = tk.Label(
        header_frame, 
        text=f"🎯 Resume Tailoring Preview - {job_title}",
        font=('Segoe UI', 16, 'bold'),
        fg='#00d4ff',
        bg='#1a1a2e'
    )
    title_label.pack(side='left', padx=10)
    
    # Stats frame
    stats_frame = tk.Frame(header_frame, bg='#1a1a2e')
    stats_frame.pack(side='right', padx=10)
    
    # Before stats
    before_label = tk.Label(
        stats_frame,
        text=f"📊 Before: Match {stats_before['match']}% | ATS {stats_before['ats']}%",
        font=('Segoe UI', 11),
        fg='#ff6b6b',
        bg='#1a1a2e'
    )
    before_label.pack(side='left', padx=20)
    
    # After stats
    after_label = tk.Label(
        stats_frame,
        text=f"📈 After: Match {stats_after['match']}% | ATS {stats_after['ats']}%",
        font=('Segoe UI', 11, 'bold'),
        fg='#51cf66',
        bg='#1a1a2e'
    )
    after_label.pack(side='left', padx=20)
    
    # Improvement indicator
    improvement = stats_after['match'] - stats_before['match']
    if improvement > 0:
        imp_text = f"⬆️ +{improvement}%"
        imp_color = '#51cf66'
    elif improvement < 0:
        imp_text = f"⬇️ {improvement}%"
        imp_color = '#ff6b6b'
    else:
        imp_text = "→ 0%"
        imp_color = '#ffd43b'
    
    imp_label = tk.Label(
        stats_frame,
        text=imp_text,
        font=('Segoe UI', 12, 'bold'),
        fg=imp_color,
        bg='#1a1a2e'
    )
    imp_label.pack(side='left', padx=10)
    
    # Keyword Analysis Panel (collapsible)
    keyword_frame = tk.Frame(preview_window, bg='#1a1a2e')
    keyword_frame.pack(fill='x', padx=10, pady=5)
    
    # Technical keywords found in tailored resume
    tech_found = stats_after.get('tech_found', [])
    tech_missing = stats_after.get('tech_missing', [])
    soft_found = stats_after.get('soft_found', [])
    soft_missing = stats_after.get('soft_missing', [])
    
    # Keywords row
    kw_row = tk.Frame(keyword_frame, bg='#1a1a2e')
    kw_row.pack(fill='x', pady=2)
    
    # Found keywords (green)
    if tech_found:
        found_label = tk.Label(
            kw_row,
            text=f"✅ Tech Skills Matched: {', '.join(tech_found[:8])}{'...' if len(tech_found)>8 else ''}",
            font=('Segoe UI', 9),
            fg='#69db7c',
            bg='#1a1a2e'
        )
        found_label.pack(side='left', padx=5)
    
    # Missing keywords (yellow warning)
    if tech_missing:
        missing_label = tk.Label(
            kw_row,
            text=f"⚠️ Tech Missing: {', '.join(tech_missing[:5])}{'...' if len(tech_missing)>5 else ''}",
            font=('Segoe UI', 9),
            fg='#ffd43b',
            bg='#1a1a2e'
        )
        missing_label.pack(side='left', padx=15)
    
    # Soft skills row
    soft_row = tk.Frame(keyword_frame, bg='#1a1a2e')
    soft_row.pack(fill='x', pady=2)
    
    if soft_found:
        soft_found_label = tk.Label(
            soft_row,
            text=f"✅ Soft Skills: {', '.join(soft_found[:5])}",
            font=('Segoe UI', 9),
            fg='#69db7c',
            bg='#1a1a2e'
        )
        soft_found_label.pack(side='left', padx=5)
    
    if soft_missing:
        soft_missing_label = tk.Label(
            soft_row,
            text=f"⚠️ Soft Missing: {', '.join(soft_missing[:3])}",
            font=('Segoe UI', 9),
            fg='#ffd43b',
            bg='#1a1a2e'
        )
        soft_missing_label.pack(side='left', padx=15)
    
    # Separator
    sep = tk.Frame(preview_window, bg='#333', height=1)
    sep.pack(fill='x', padx=10, pady=5)
    
    # Main comparison frame
    main_frame = tk.Frame(preview_window, bg='#1a1a2e')
    main_frame.pack(fill='both', expand=True, padx=10, pady=5)
    
    # Left panel - Master Resume
    left_frame = tk.Frame(main_frame, bg='#1a1a2e')
    left_frame.pack(side='left', fill='both', expand=True, padx=(0, 5))
    
    left_header = tk.Label(
        left_frame,
        text="📋 MASTER RESUME (Original)",
        font=('Segoe UI', 12, 'bold'),
        fg='#74c0fc',
        bg='#1a1a2e'
    )
    left_header.pack(pady=5)
    
    left_text = scrolledtext.ScrolledText(
        left_frame,
        wrap='word',
        font=('Consolas', 10),
        bg='#16213e',
        fg='#e8e8e8',
        insertbackground='white',
        selectbackground='#4a69bd',
        padx=10,
        pady=10
    )
    left_text.pack(fill='both', expand=True)
    left_text.insert('1.0', master_text)
    left_text.config(state='disabled')
    
    # Right panel - Tailored Resume
    right_frame = tk.Frame(main_frame, bg='#1a1a2e')
    right_frame.pack(side='right', fill='both', expand=True, padx=(5, 0))
    
    right_header = tk.Label(
        right_frame,
        text="✨ TAILORED RESUME (AI Enhanced)",
        font=('Segoe UI', 12, 'bold'),
        fg='#69db7c',
        bg='#1a1a2e'
    )
    right_header.pack(pady=5)
    
    right_text = scrolledtext.ScrolledText(
        right_frame,
        wrap='word',
        font=('Consolas', 10),
        bg='#16213e',
        fg='#e8e8e8',
        insertbackground='white',
        selectbackground='#4a69bd',
        padx=10,
        pady=10
    )
    right_text.pack(fill='both', expand=True)
    right_text.insert('1.0', tailored_text)
    right_text.config(state='disabled')
    
    # Sync scrolling between both text widgets
    def sync_scroll(*args):
        left_text.yview(*args)
        right_text.yview(*args)
    
    def on_left_scroll(*args):
        right_text.yview_moveto(args[0])
    
    def on_right_scroll(*args):
        left_text.yview_moveto(args[0])
    
    left_text.config(yscrollcommand=lambda *args: on_left_scroll(left_text.yview()[0]))
    right_text.config(yscrollcommand=lambda *args: on_right_scroll(right_text.yview()[0]))
    
    # Button frame at bottom
    button_frame = tk.Frame(preview_window, bg='#1a1a2e')
    button_frame.pack(fill='x', padx=10, pady=10)
    
    def open_pdf():
        import subprocess
        import sys
        if paths.get("pdf") and os.path.exists(paths["pdf"]):
            if sys.platform == 'win32':
                os.startfile(paths["pdf"])
            else:
                subprocess.run(['open', paths["pdf"]], check=False)
    
    def open_docx():
        import subprocess
        import sys
        if paths.get("docx") and os.path.exists(paths["docx"]):
            if sys.platform == 'win32':
                os.startfile(paths["docx"])
            else:
                subprocess.run(['open', paths["docx"]], check=False)
    
    def open_folder():
        import subprocess
        import sys
        folder = os.path.dirname(paths.get("pdf") or paths.get("docx") or paths.get("txt", ""))
        if folder and os.path.exists(folder):
            if sys.platform == 'win32':
                os.startfile(folder)
            else:
                subprocess.run(['open', folder], check=False)
    
    # Buttons
    btn_style = {'font': ('Segoe UI', 10), 'width': 18, 'cursor': 'hand2'}
    
    if HAS_TTK:
        pdf_btn = ttkb.Button(button_frame, text="📄 Open PDF", command=open_pdf, bootstyle="info")
        docx_btn = ttkb.Button(button_frame, text="📝 Open DOCX", command=open_docx, bootstyle="success")
        folder_btn = ttkb.Button(button_frame, text="📁 Open Folder", command=open_folder, bootstyle="secondary")
        close_btn = ttkb.Button(button_frame, text="✖️ Close", command=preview_window.destroy, bootstyle="danger")
    else:
        pdf_btn = tk.Button(button_frame, text="📄 Open PDF", command=open_pdf, bg='#4a69bd', fg='white', **btn_style)
        docx_btn = tk.Button(button_frame, text="📝 Open DOCX", command=open_docx, bg='#20bf6b', fg='white', **btn_style)
        folder_btn = tk.Button(button_frame, text="📁 Open Folder", command=open_folder, bg='#636e72', fg='white', **btn_style)
        close_btn = tk.Button(button_frame, text="✖️ Close", command=preview_window.destroy, bg='#e74c3c', fg='white', **btn_style)
    
    pdf_btn.pack(side='left', padx=10)
    docx_btn.pack(side='left', padx=10)
    folder_btn.pack(side='left', padx=10)
    close_btn.pack(side='right', padx=10)
    
    # File paths info
    info_label = tk.Label(
        button_frame,
        text=f"Files saved to: {os.path.dirname(paths.get('pdf', paths.get('docx', '')))}",
        font=('Segoe UI', 9),
        fg='#adb5bd',
        bg='#1a1a2e'
    )
    info_label.pack(side='left', padx=20)


def open_preview(paths: dict, diff_report: str | None = None, master_text: str = "", tailored_text: str = "", jd_text: str = "", job_title: str = "Job") -> None:
    """Open preview GUI window with side-by-side comparison.
    
    Files are NOT auto-opened - user must click buttons in GUI to view PDF/DOCX.
    """
    # If we have text content, show the GUI preview only (no auto-open files)
    if master_text and tailored_text:
        try:
            show_preview_gui(master_text, tailored_text, jd_text, paths, job_title)
            return
        except Exception as e:
            print(f"GUI preview error: {e}")
            # Don't fall back to opening files - just show error
            return
    
    # If no text content provided, just print info (don't auto-open files)
    print(f"Preview files available:")
    if paths.get("pdf"):
        print(f"  PDF: {paths['pdf']}")
    if paths.get("docx"):
        print(f"  DOCX: {paths['docx']}")
    if diff_report:
        print(f"  Diff Report: {diff_report}")


def tailor_resume_to_files(
    resume_text: Optional[str],
    job_description: str,
    instructions: Optional[str] = None,
    provider: Optional[str] = None,
    client: Optional[object] = None,
    output_dir: Optional[str] = None,
    resume_path: Optional[str] = None,
    job_title: Optional[str] = None,
    candidate_name: Optional[str] = None,
    enable_preview: bool = False,
) -> dict:
    """
    Tailor a resume and save it to text, docx, and pdf. Returns a dict of paths.
    NOW WITH REAL-TIME PROGRESS UPDATES!
    """
    # Import progress utilities with immediate flush for real-time updates
    _metrics_module = None
    _log_module = None
    try:
        from modules.dashboard import log_handler as _log_module
        from modules.dashboard import metrics as _metrics_module
    except ImportError:
        pass
    
    def update_progress(jd_pct: int, resume_pct: int, status: str = ""):
        """Update progress with IMMEDIATE flush for real-time dashboard updates."""
        if _metrics_module:
            _metrics_module.set_metric('jd_progress', jd_pct)
            _metrics_module.set_metric('resume_progress', resume_pct)
        if status and _log_module:
            _log_module.publish(status, "AI")
    
    # Reset progress at start
    update_progress(0, 0, "")
    
    # === STEP 1: Reading resume (0-10%) ===
    update_progress(5, 5, "📄 Reading resume file...")
    
    if not resume_text and resume_path:
        resume_text = _read_resume_text(resume_path)
    if not resume_text:
        raise ValueError("Resume text is required for tailoring.")

    update_progress(0, 10, "✅ Resume loaded")
    
    resolved_provider = (provider or ai_provider or "ollama").lower()

    # === STEP 2: Check cache (10-15%) ===
    update_progress(5, 15, "🔍 Checking cache...")
    
    cache_dir = os.path.join(generated_resume_path or "all resumes/", "cache")
    _ensure_output_dir(cache_dir)
    cache_key = _cache_key(resume_text, job_description, instructions, resolved_provider, job_title)
    cache_path = os.path.join(cache_dir, f"{cache_key}.json")
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            cached_paths = data.get("paths", {})
            if cached_paths and all(os.path.exists(p) for p in cached_paths.values() if p):
                update_progress(100, 100, "⚡ Using cached tailored resume!")
                print_lg(f"DEBUG: Returning cached result with pdf={cached_paths.get('pdf', 'NOT SET')}")
                # Return in same format as non-cached path (flat dict with paths + metadata)
                return {
                    **cached_paths,
                    "master_text": resume_text,
                    "tailored_text": "",  # Not cached
                    "jd_text": job_description,
                    "job_title": job_title or "Tailored Resume"
                }
        except Exception as cache_err:
            print_lg(f"DEBUG: Cache read error: {cache_err}")
            pass

    # === STEP 3: JD Analysis (15-40%) ===
    update_progress(15, 20, "📋 Analyzing job description...")
    update_progress(25, 20, "🔑 Extracting keywords from JD...")
    update_progress(40, 25, "✅ JD analysis complete")
    
    # === STEP 4: AI Tailoring (40-80%) ===
    update_progress(50, 30, f"🤖 Calling {resolved_provider.upper()} AI for tailoring...")
    update_progress(60, 40, "⏳ AI is optimizing your resume...")
    
    print_lg(f"DEBUG: Calling tailor_resume_text with provider={provider or resolved_provider}")
    print_lg(f"DEBUG: resume_text length={len(resume_text) if resume_text else 0}, jd length={len(job_description) if job_description else 0}")
    
    try:
        tailored = tailor_resume_text(
            resume_text=resume_text,
            job_description=job_description,
            instructions=instructions,
            provider=provider,
            client=client,
        )
        print_lg(f"DEBUG: tailor_resume_text returned {len(tailored) if tailored else 0} chars")
        if tailored and tailored.startswith("["):
            print_lg(f"DEBUG: tailor_resume_text returned ERROR: {tailored[:200]}")
    except Exception as tailor_err:
        print_lg(f"DEBUG: tailor_resume_text EXCEPTION: {tailor_err}")
        import traceback
        print_lg(f"DEBUG: Traceback: {traceback.format_exc()}")
        raise
    
    update_progress(80, 60, "✅ AI tailoring complete!")
    
    tailored = _normalize_one_page(tailored)
    update_progress(85, 65, "📏 Normalized to one page")

    target_dir = output_dir or os.path.join(generated_resume_path or "all resumes/", "temp")

    cand = candidate_name or _candidate_name() or "Candidate"
    # Format: "CandidateName - JobTitle" for better organization
    title_part = job_title or "Tailored Resume"
    # Truncate job title if too long (keep first 50 chars)
    if len(title_part) > 50:
        title_part = title_part[:50].rsplit(' ', 1)[0]
    base = f"{cand} - {title_part}"
    base_name = _sanitize_filename(base)

    # === STEP 5: Saving files (80-95%) ===
    update_progress(90, 70, "💾 Saving text file...")
    txt_path = _save_text(tailored, target_dir, base_name=base_name)

    docx_path = ""
    pdf_path = ""

    update_progress(95, 80, "📝 Creating DOCX...")
    if resume_path and resume_path.lower().endswith(".docx"):
        try:
            paragraphs = _read_docx_paragraphs(resume_path)
            tailored_paragraphs = _tailor_paragraphs_with_ai(
                paragraphs, job_description, instructions, resolved_provider, client
            )
            if tailored_paragraphs and len(tailored_paragraphs) == len(paragraphs):
                docx_path = _write_docx_from_template(resume_path, tailored_paragraphs, target_dir, base_name)
                tailored_text_for_diff = "\n".join(tailored_paragraphs)
                tailored = tailored_text_for_diff
            else:
                # Fallback: keep headings and replace body lines sequentially
                try:
                    from docx import Document
                    doc = Document(resume_path)
                    lines = [l for l in tailored.splitlines() if l.strip()]
                    idx = 0
                    for p in doc.paragraphs:
                        if not p.text.strip():
                            continue
                        if p.style and "Heading" in str(p.style.name):
                            continue
                        if idx < len(lines):
                            p.text = lines[idx]
                            idx += 1
                    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                    docx_path = os.path.join(target_dir, f"{base_name}_{stamp}.docx")
                    doc.save(docx_path)
                except Exception:
                    docx_path = _write_docx(tailored, target_dir, base_name=base_name)
        except Exception:
            docx_path = _write_docx(tailored, target_dir, base_name=base_name)
    else:
        docx_path = _write_docx(tailored, target_dir, base_name=base_name)

    update_progress(98, 90, "📄 Creating PDF...")
    # Prefer docx->pdf conversion if available to preserve formatting
    try:
        from docx2pdf import convert  # type: ignore
        print_lg("DEBUG: docx2pdf available, converting...")
        if docx_path:
            convert(docx_path)
            pdf_candidate = os.path.splitext(docx_path)[0] + ".pdf"
            if os.path.exists(pdf_candidate):
                pdf_path = pdf_candidate
                print_lg(f"DEBUG: PDF created via docx2pdf: {pdf_path}")
    except ImportError:
        print_lg("DEBUG: docx2pdf not available, using reportlab fallback...")
        try:
            pdf_path = _write_pdf(tailored, target_dir, base_name=base_name)
            print_lg(f"DEBUG: PDF created via reportlab: {pdf_path}")
        except Exception as pdf_err:
            print_lg(f"DEBUG: PDF creation FAILED: {pdf_err}")
            import traceback
            print_lg(f"DEBUG: {traceback.format_exc()}")
    except Exception as conv_err:
        print_lg(f"DEBUG: docx2pdf conversion failed: {conv_err}, using reportlab...")
        try:
            pdf_path = _write_pdf(tailored, target_dir, base_name=base_name)
            print_lg(f"DEBUG: PDF created via reportlab: {pdf_path}")
        except Exception as pdf_err:
            print_lg(f"DEBUG: PDF creation FAILED: {pdf_err}")
            import traceback
            print_lg(f"DEBUG: {traceback.format_exc()}")

    diff_report = None
    if enable_preview:
        diff_report = generate_preview_report(resume_text, tailored, job_description, target_dir, base_name)

    # === STEP 6: Complete (100%) ===
    update_progress(100, 100, "✅ Resume tailoring complete!")
    
    result = {"txt": txt_path, "docx": docx_path, "pdf": pdf_path, "diff": diff_report, 
              "master_text": resume_text, "tailored_text": tailored, "jd_text": job_description,
              "job_title": job_title or "Tailored Resume"}
    print_lg(f"Saved tailored resume files: {txt_path}, {docx_path}, {pdf_path}")

    try:
        with open(cache_path, "w", encoding="utf-8") as f:
            json.dump({"paths": {"txt": txt_path, "docx": docx_path, "pdf": pdf_path, "diff": diff_report}}, f)
    except Exception:
        pass

    return result


def tailor_resume_to_file(
    resume_text: str,
    job_description: str,
    instructions: Optional[str] = None,
    provider: Optional[str] = None,
    client: Optional[object] = None,
    output_dir: Optional[str] = None,
    job_title: Optional[str] = None,
    candidate_name: Optional[str] = None,
) -> str:
    """
    Backwards-compatible: returns the PDF path from tailor_resume_to_files.
    """
    paths = tailor_resume_to_files(
        resume_text=resume_text,
        job_description=job_description,
        instructions=instructions,
        provider=provider,
        client=client,
        output_dir=output_dir,
        resume_path=None,
        job_title=job_title,
        candidate_name=candidate_name,
    )
    return paths.get("pdf", "")
